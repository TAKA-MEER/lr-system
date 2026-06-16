import logging
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ---- スタイル定数 ----
COLOR_HEADER = RGBColor(0x1F, 0x49, 0x7D)   # 濃い青
COLOR_OUR    = RGBColor(0x00, 0x70, 0xC0)   # 自社ラベル色
COLOR_CLIENT = RGBColor(0xC0, 0x50, 0x00)   # 相手方ラベル色


def generate_docx(minutes: dict, output_path: str):
    """
    minutes JSON を受け取り、一般的な議事録フォーマットの docx を生成する。
    社内テンプレートが決まり次第、この関数を差し替える。
    """
    doc = Document()

    # ページ余白を設定
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ---- タイトル ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("立 会 試 験  議 事 録")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_HEADER

    doc.add_paragraph()  # 空行

    # ---- 基本情報テーブル ----
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    _set_col_width(info_table, 0, Cm(3.5))
    _set_col_width(info_table, 1, Cm(13.0))

    attendees = minutes.get("attendees", {})
    client_names  = "、".join(attendees.get("client", []))  or "（未入力）"
    our_names     = "、".join(attendees.get("our_side", [])) or "（未入力）"

    info_rows = [
        ("試験名",   minutes.get("trial_name", "")),
        ("日時",     minutes.get("date", "")),
        ("場所",     minutes.get("location", "")),
        ("参加者",   f"相手方: {client_names}\n自社: {our_names}"),
    ]
    for i, (label, value) in enumerate(info_rows):
        row = info_table.rows[i]
        _set_cell(row.cells[0], label, bold=True, bg_color="D9E1F2")
        _set_cell(row.cells[1], value)

    doc.add_paragraph()

    # ---- 協議事項 ----
    _add_section_heading(doc, "１．協議事項")

    discussions = minutes.get("discussions", [])
    if discussions:
        d_table = doc.add_table(rows=1, cols=5)
        d_table.style = "Table Grid"
        _set_col_width(d_table, 0, Cm(0.8))
        _set_col_width(d_table, 1, Cm(3.5))
        _set_col_width(d_table, 2, Cm(4.5))
        _set_col_width(d_table, 3, Cm(4.5))
        _set_col_width(d_table, 4, Cm(2.0))

        headers = ["No", "議題", "相手方 要望・確認事項", "自社 回答・対応", "状態"]
        for j, h in enumerate(headers):
            _set_cell(d_table.rows[0].cells[j], h, bold=True, bg_color="D9E1F2", center=True)

        for idx, d in enumerate(discussions, 1):
            row = d_table.add_row()
            _set_cell(row.cells[0], str(idx), center=True)
            _set_cell(row.cells[1], d.get("topic", ""))
            _set_cell(row.cells[2], d.get("client_request") or "")
            _set_cell(row.cells[3], d.get("our_response") or "")
            status = d.get("status", "")
            _set_cell(row.cells[4], status, center=True)
    else:
        doc.add_paragraph("（協議事項なし）")

    doc.add_paragraph()

    # ---- アクションアイテム ----
    _add_section_heading(doc, "２．アクションアイテム")

    actions = minutes.get("action_items", [])
    if actions:
        a_table = doc.add_table(rows=1, cols=4)
        a_table.style = "Table Grid"
        _set_col_width(a_table, 0, Cm(0.8))
        _set_col_width(a_table, 1, Cm(8.5))
        _set_col_width(a_table, 2, Cm(2.5))
        _set_col_width(a_table, 3, Cm(3.0))

        for j, h in enumerate(["No", "対応内容", "担当", "期限"]):
            _set_cell(a_table.rows[0].cells[j], h, bold=True, bg_color="D9E1F2", center=True)

        for idx, a in enumerate(actions, 1):
            row = a_table.add_row()
            owner_label = "自社" if a.get("owner") == "our_side" else "相手方" if a.get("owner") == "client" else a.get("owner", "")
            _set_cell(row.cells[0], str(idx), center=True)
            _set_cell(row.cells[1], a.get("content", ""))
            _set_cell(row.cells[2], owner_label, center=True)
            _set_cell(row.cells[3], a.get("deadline") or "", center=True)
    else:
        doc.add_paragraph("（アクションアイテムなし）")

    doc.add_paragraph()

    # ---- 備考 ----
    _add_section_heading(doc, "３．備考")
    doc.add_paragraph(minutes.get("notes", ""))
    doc.add_paragraph()

    # ---- 承認欄 ----
    _add_section_heading(doc, "４．承認")
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"
    _set_col_width(sig_table, 0, Cm(8.0))
    _set_col_width(sig_table, 1, Cm(8.0))
    _set_cell(sig_table.rows[0].cells[0], "相手方 確認", bold=True, bg_color="D9E1F2", center=True)
    _set_cell(sig_table.rows[0].cells[1], "自社 確認",   bold=True, bg_color="D9E1F2", center=True)
    _set_cell(sig_table.rows[1].cells[0], "\n\n")
    _set_cell(sig_table.rows[1].cells[1], "\n\n")

    doc.save(output_path)
    logger.info(f"docx 生成完了: {output_path}")


# ------------------------------------------------------------------
# ヘルパー関数
# ------------------------------------------------------------------

def _add_section_heading(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_HEADER


def _set_cell(cell, text: str, bold: bool = False, bg_color: str | None = None, center: bool = False):
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)

    if bg_color:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_color)
        tc_pr.append(shd)


def _set_col_width(table, col_idx: int, width):
    for row in table.rows:
        row.cells[col_idx].width = width
